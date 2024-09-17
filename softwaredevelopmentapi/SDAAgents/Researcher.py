from huggingface_hub import login
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import json
import textwrap
import os
#####################
import os
import shutil
import torch
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline
import torch
from docx import Document
from softwaredevelopmentapi.model_loader import llm_model,tokenizer
class ResearcherAgent:
    def __init__(self, model_name="facebook/llama-2", tokenizer_name="facebook/llama-2"):
        #self.tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True,
        #                                               use_auth_token="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")
        #self.model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,
        #                                                  use_auth_token="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")
        self.tokenizer = tokenizer
        self.model = llm_model
    def read_pdf(self, file_path):
        loader = PyPDFLoader(file_path)
        data = loader.load()
            
        return data

    def read_docx(self, file_path):
        loader = Docx2txtLoader(file_path)
        data = loader.load()
            
        return data

    def read_txt(self, file_path):
        loader = TextLoader(file_path)
        data = loader.load()
            
        return data
    def add_text_to_docx(self, file_path, text, output_path):
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(output_path)
        
    def generate_response_from_document(self, file_path, query, project_info):
        if len(project_info) > 0:
            self.add_text_to_docx(file_path, project_info, "Knowldgebase/temp/temp.docx")
        """Generate a response based on a user query from a document."""
        # Determine the file type and extract text
        if file_path.lower().endswith('.pdf'):
            document_text = self.read_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            document_text = self.read_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            document_text = self.read_txt(file_path)
        else:
            raise ValueError("Unsupported file type. Supported types are: 'pdf', 'docx', 'txt'.")

        streamer = TextStreamer(self.tokenizer, skip_prompt=True, skip_special_tokens=True)

        text_pipeline = pipeline(
            "text-generation",
            model=self.model,
            tokenizer=self.tokenizer,
            max_new_tokens=1024,
            temperature=0.5,
            top_p=0.95,
            repetition_penalty=1.15,
            streamer=streamer,
        )

        # document_text[0]= f"\n{document_text[0]=}\n{project_info}"

        llm = HuggingFacePipeline(pipeline=text_pipeline, model_kwargs={"temperature": 0.5})

        # Step 2: Split the text into manageable chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
        texts = text_splitter.split_documents(document_text)

        # Step 3: Generate embeddings for the texts
        embeddings = HuggingFaceEmbeddings(
            model_name="hkunlp/instructor-xl", model_kwargs={"device": "cuda"}
        )
        db = Chroma.from_documents(texts, embeddings)

        #         SYSTEM_PROMPT = """
        #         You are a knowledgeable and precise assistant designed to handle document-related queries efficiently and safely. Always respond as accurately as possible, ensuring your answers comply with ethical guidelines and legal standards. Avoid sharing any information that could be harmful, unethical, biased, or illegal.

        #         If a query is ambiguous or unclear, ask clarifying questions to better understand the user's needs before providing a response. If the system cannot retrieve the requested information due to format limitations or content availability, inform the user clearly and suggest possible alternatives.

        #         Ensure that your responses enhance user interaction with the system, guiding them on how to upload, retrieve, and manage their documents effectively. If you encounter a question outside your expertise or capability, advise the user accordingly without resorting to misleading or incorrect information.
        #         """
        SYSTEM_PROMPT = f"""
       You are a knowledgeable Software Developer Assistant, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with software engineering standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant to provided project in the document.
Always clarify the context if needed. the query is {query}
        """
        prompt = PromptTemplate(template=""" You are a knowledgeable Software Developer Assistant, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with software engineering standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant to provided project in the document.
Always clarify the context if needed.

    Context: {context}
    User: {question}
    Software Developer Assistant:""", input_variables=["context", "question"])

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=db.as_retriever(search_kwargs={"k": 2}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt},
        )

        # Execute a query (you can customize this part as needed)
        query_result = qa_chain(SYSTEM_PROMPT)

        # Execute the query
        #         query_result, source_documents = qa_chain.run(query)

        # print("llm!!!!!!!!!!!!!!!!!!!!! Result",query_result)

        # Extract the answer from the 'response' field using string manipulation
        # Assuming the answer always follows "Answer:" and ends at the end of the string
        response_content = query_result['result']
        answer_prefix = "Software Developer Assistant: "
        answer_start_index = response_content.find(answer_prefix)
        if answer_start_index != -1:
            answer = response_content[answer_start_index + len(answer_prefix):].strip()
            print(answer)
            return answer
        else:
            print("No answer found in the response.")
            return response_content

    def researcher_agent(self, file_path, query,project_info):
        sections_queries = {
            "title": "Provide a suitable title for the research based on the project description.",
            "abstract": "Write an abstract summarizing the main points and findings of the research.",
            "introduction": "Provide an introduction that includes the background and significance of the research.",
            "literature review": "Conduct a literature review highlighting the key studies and findings relevant to the research topic.",
            "methodology": "Describe the research methodology, including the research design, data collection, and analysis methods.",
            "results": "Present the research results, including any data and findings.",
            "discussion": "Discuss the implications of the research results, including interpretation and analysis.",
            "conclusion": "Provide a conclusion summarizing the key findings and their significance.",
            "references": "List all references cited in the research.",
            "appendices": "Include any additional information or data in the appendices."
        }

        
        for section, query in sections_queries.items():
            response = self.generate_response_from_document(file_path, query, project_info)
            
            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            
            # Append the response to the project_info with a newline
            if project_info:
                project_info += f"\n{section_key}:\n{response}"
            else:
                project_info = f"{section_key}:\n{response}"

        # Create a new .docx document
        doc = DocxDocument()
        doc.add_heading('Research Paper', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])
        
        # Save the document
        doc.save(output_path)
        return "Researcg has been completed."
