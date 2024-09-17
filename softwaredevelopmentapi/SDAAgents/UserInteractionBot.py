import json
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
from docx import Document
import textwrap
from softwaredevelopmentapi.model_loader import llm_model,tokenizer
from docx import Document
import os

class ConversationSaver:
    def __init__(self, file_path):
        self.file_path = file_path
        if os.path.exists(file_path):
            self.conversation = Document(file_path)  # Load existing document
        else:
            self.conversation = Document()  # Create a new document

    def save_conversation(self, user_query, bot_response):
        self.conversation.add_heading('Bot Query:', level=2)
        self.conversation.add_paragraph(user_query)
        self.conversation.add_heading('User Response:', level=2)
        self.conversation.add_paragraph(bot_response)

    def save_to_docx(self):
        self.conversation.save(self.file_path)
        
class UserInteractionBot:
    def __init__(self, model_name="meta-llama/Llama-2-7b-chat-hf", tokenizer_name="meta-llama/Llama-2-7b-chat-hf"):
        #self.device = "cuda" if torch.cuda.is_available() else "cpu"
        #self.model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,use_auth_token="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")
        self.model = llm_model
        #self.tokenizer = AutoTokenizer.from_pretrained(tokenizer_name,use_fast=True,device=self.device,use_auth_token="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")
        self.tokenizer =tokenizer
        self.terminators = [
            tokenizer.eos_token_id,
            tokenizer.convert_tokens_to_ids("<|eot_id|>")
        ]
        self.conversation = Document()

    def get_prompt(self, instruction, system_prompt):
        prompt_template = f"{system_prompt}, {instruction}"
        return prompt_template

    def cut_off_text(self, text, prompt):
        cutoff_phrase = prompt
        index = text.find(cutoff_phrase)
        if index != -1:
            return text[:index]
        else:
            return text

    def remove_substring(self, string, substring):
        return string.replace(substring, "")

    def generate(self, text, system_prompt, max_new_tokens):
        prompt = self.get_prompt(text, system_prompt)
        with torch.autocast('cuda', dtype=torch.bfloat16):
            inputs = self.tokenizer(prompt, return_tensors="pt").to('cuda')
            outputs = self.model.generate(
                **inputs,
                max_length=100,  # Adjust max_length as needed
                max_new_tokens=100,  # Increase max_new_tokens
                eos_token_id=self.tokenizer.eos_token_id,
                pad_token_id=self.tokenizer.eos_token_id
            )
            final_outputs = self.tokenizer.batch_decode(outputs, skip_special_tokens=True)[0]
            final_outputs = self.cut_off_text(final_outputs, '</s>')
            final_outputs = self.remove_substring(final_outputs, prompt)
        return final_outputs

    

    def parse_text(self, text):
        wrapped_text = textwrap.fill(text, width=100)
        return wrapped_text + '\n\n'

    def gen_chat(self, query):
        DEFAULT_SYSTEM_PROMPT = """You are an AI trained to gather user requirements for website development. Ask provided question . Only Ask question."""

         # """You are an AI trained to gather user requirements for website development. Question is provided below Ask provided question in interactive way to understand the project. Only Ask question."""

        prompt = f"Question is '{query}'"

        generated_text = self.generate(prompt, DEFAULT_SYSTEM_PROMPT, 100)
        response = self.parse_text(generated_text)

        #self.save_conversation(query, response)
        return response

    def ask_question_with_llm(self, system_prompt, instruction):
        prompt = self.get_prompt(instruction, system_prompt)
        question = self.generate(prompt, system_prompt, max_new_tokens=50)
        print(f"Question: {question}")
        #user_response = input("Your response: ")
        return question

    def get_question_by_key(self,key,questions):
        if key in questions:
            return questions[key]
        else:
            return "Invalid key: key should be between 1 and {}".format(len(questions))
    def run_bot(self,user_response,key, output_path,projectid):
        key=int(key)
        questions = {
            1:"What is the main purpose of your website?",
            2:"Who are the primary users of your website?",
            3:"What key features should your website include?",
            4:"Do you have any specific design or aesthetic preferences?",
            5:"What type of content will your website host (e.g., articles, videos, products)?",
            6:"Are there any technical constraints or requirements for your website?",
            7:"Do you require backend services like databases or user management systems?",
            8:"What is your budget for this project?",
            9:"What is your desired timeline for launching the website?",
            10:"How do you plan to handle website maintenance and updates post-launch?",
            11:"Do you anticipate any future expansions or feature additions?",
            12:"Are there any specific legal or compliance requirements your website must adhere to?"
        }

        system_prompt = "You are an AI trained to gather user requirements for website development. Ask question to understand the project in detail."
        #for question in questions:
        #    self.ask_question_with_llm(system_prompt, question)
        #for key, question in questions.items():
        #    query = f"For Question {key}, {question}"
        #    response = self.generate_response_from_document(file_path, query, project_info)
        if key ==0:
            heading = "User Requirements"
            file_path = f'Knowldgebase/client/conversation{projectid}.docx'
            saver = ConversationSaver(file_path)
            saver.save_conversation(heading, user_response)
            saver.save_to_docx()
            key = key + 1
            question = self.get_question_by_key(key, questions)
            response = self.ask_question_with_llm(system_prompt, question)
            print(f"Conversation saved to {file_path}")
            return response, key
        elif  1 <= key <= len(questions):
            question = self.get_question_by_key(key, questions)
            file_path = f'Knowldgebase/client/conversation{projectid}.docx'
            saver = ConversationSaver(file_path)
            saver.save_conversation(question, user_response)
            saver.save_to_docx()
            key = key + 1
            question = self.get_question_by_key(key, questions)
            response = self.ask_question_with_llm(system_prompt, question)
            print(f"Conversation saved to {file_path}")
            return response, key
        elif key == len(questions):
            question = "Software Requirements has been gathered."
            response = self.ask_question_with_llm(system_prompt, question)
            #file_path = f'Knowldgebase/client/conversation{projectid}.docx'
            #saver = ConversationSaver(file_path)
            #saver.save_conversation(heading, user_response)
            #saver.save_to_docx()
            return response, key