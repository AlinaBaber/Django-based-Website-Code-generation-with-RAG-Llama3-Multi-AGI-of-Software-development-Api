import os
import tiktoken
from langchain.chains import RetrievalQA, ConversationalRetrievalChain
from langchain_openai import ChatOpenAI
from langchain.document_loaders import TextLoader
from langchain.vectorstores import FAISS
from langchain.memory import ConversationBufferMemory
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.schema import Document
import docx
import docx
from langchain.schema import Document as LangChainDocument  # Rename for clarity

class OpenAIChat:
    def __init__(self, api_key: str, model_name: str = "gpt-4", temperature: float = 0.7):
        os.environ["OPENAI_API_KEY"] = api_key
        self.llm = ChatOpenAI(temperature=temperature, model_name=model_name)
        self.memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
        self.vectorstore = None

    def load_docx(self, file_path: str):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        self.data = "\n".join(full_text)

    def split_text(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        text_splitter = CharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        document = Document(page_content=self.data)
        self.data = text_splitter.split_documents([document])

    def create_vectorstore(self):
        embeddings = OpenAIEmbeddings()
        self.vectorstore = FAISS.from_documents(self.data, embedding=embeddings)

    def create_conversation_chain(self):
        self.conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=self.llm,
            chain_type="stuff",
            retriever=self.vectorstore.as_retriever(),
            memory=self.memory
        )

    def ask_question(self, query: str) -> str:
        result = self.conversation_chain({"question": query})
        return result["answer"]

    # Rest of your class code remains the same...
    
    # Adjust your function to explicitly use docx.Document
    def store_response_in_docx(self,response_text, output_path):
        doc = docx.Document()  # Create a new Document using docx
        doc.add_paragraph(response_text)  # Add the generated response as a paragraph
        doc.save(output_path)  # Save the document to the specified path
    
def generatecode_openai(docx_file_path,output_doc_path, projectid,model_name="gpt-4"):
    # Usage example
    # Prompt the user for their OpenAI API key
    api_key ="put api key here"
    
    # Initialize the OpenAIChat class
    chat = OpenAIChat(api_key=api_key, model_name=model_name)
    
    # Load DOCX data
    docx_file_path = docx_file_path
    chat.load_docx(file_path=docx_file_path)
    # Split text into chunks
    chat.split_text()
    
    # Create vector store
    chat.create_vectorstore()
    
    # Create conversation chain
    chat.create_conversation_chain()
    # Ask questions
    # Generate Django models from SRS
    query = """
    I am working on a project that requires the development of a complete static website based on a Software Requirements Specification (SRS) document. The SRS outlines the necessary features, functionality, and design elements for the website. To begin the development process, I need to identify all the HTML pages that need to be created.
    
    Please provide a detailed list of HTML pages that are required according to the SRS. For each page, include a brief description of its purpose, key features, and any specific elements or sections that should be included.:
    """ + chat.data[0].page_content
    code = chat.ask_question(query)
    print(code)
    
    # Store the generated Django models in a DOCX file
    output_doc_path=f'Knowldgebase/SA/ListofHTMLPage{projectid}.docx'
    chat.store_response_in_docx(code, output_doc_path)
    
    # Load DOCX data
    docx_file_path = f'Knowldgebase/SA/ListofHTMLPage{projectid}.docx'
    chat.load_docx(file_path=docx_file_path)
    # Split text into chunks
    chat.split_text()
    
    # Create vector store
    chat.create_vectorstore()
    
    # Create conversation chain
    chat.create_conversation_chain()
    # Ask questions
    # Generate Django models from SRS
    query = """
    Generate the complete HTML pages, CSS, and JavaScript code for a static website based on the provided document content. The website should include:
    Add sections and HTML Pages  according to your own knowldged about similar project.
    Ensure the HTML, CSS, and JavaScript are in separate files. Provide the complete code for each file. Here is the SRS document content:
    """ + chat.data[0].page_content
    code = chat.ask_question(query)
    print(code)
    # Store the generated Django models in a DOCX file
    #output_doc_path=f'Knowldgebase/SD/HTMLS{projectid}.docx'
    chat.store_response_in_docx(code, output_doc_path)
    return "Code has been generated."