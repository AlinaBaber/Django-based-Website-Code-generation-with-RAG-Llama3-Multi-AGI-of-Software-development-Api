from huggingface_hub import login
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import json
import textwrap
import os
#####################
import os
import shutil
import torch
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline
import torch
from docx import Document
import os
import json
import os
import re

class SoftwareDeveloperAgent:
    def __init__(self, model_name="facebook/llama-2", tokenizer_name="facebook/llama-2"):
#        self.tokenizer = AutoTokenizer.from_pretrained("codellama/CodeLlama-7b-hf", use_fast=True, use_auth_token="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")
#        self.model = AutoModelForCausalLM.from_pretrained("codellama/CodeLlama-7b-hf", device_map='auto', torch_dtype=torch.float16, use_auth_token="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")
        self.tokenizer = AutoTokenizer.from_pretrained("bigcode/starcoder2-7b", use_fast=True)
        self.model = AutoModelForCausalLM.from_pretrained("bigcode/starcoder2-7b", device_map='auto', torch_dtype=torch.float16)    
    def read_pdf(self, file_path):
        loader = PyPDFLoader(file_path)
        data = loader.load()
            
        return data

    def read_docx(self, file_path):
        loader = Docx2txtLoader(file_path)
        data = loader.load()
            
        return data

    def read_txt(self, file_path):
        loader = TextLoader(file_path)
        data = loader.load()
            
        return data

    def add_text_to_docx(self, file_path, text, output_path):
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(output_path)
        
    def generate_response_from_document(self, file_path, query, project_info):
        if len(project_info)>0:
            self.add_text_to_docx(file_path, project_info, "Knowldgebase/temp/temp.docx")
        """Generate a response based on a user query from a document."""
        # Determine the file type and extract text
        if file_path.lower().endswith('.pdf'):
            document_text = self.read_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            document_text = self.read_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            document_text = self.read_txt(file_path)
        else:
            raise ValueError("Unsupported file type. Supported types are: 'pdf', 'docx', 'txt'.")
            
        streamer = TextStreamer(self.tokenizer, skip_prompt=True, skip_special_tokens=True)

        text_pipeline = pipeline(
            "text-generation",
            model=self.model,
            tokenizer=self.tokenizer,
            max_new_tokens=1024,
            temperature=0.5,
            top_p=0.95,
            repetition_penalty=1.15,
            streamer=streamer,
        )

       # document_text[0]= f"\n{document_text[0]=}\n{project_info}"

        llm = HuggingFacePipeline(pipeline=text_pipeline, model_kwargs={"temperature": 0.5})

        # Step 2: Split the text into manageable chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
        texts = text_splitter.split_documents(document_text)

        # Step 3: Generate embeddings for the texts
        embeddings = HuggingFaceEmbeddings(
            model_name="hkunlp/instructor-xl", model_kwargs={"device": "cuda"}
        )
        db = Chroma.from_documents(texts, embeddings)


#         SYSTEM_PROMPT = """
#         You are a knowledgeable and precise assistant designed to handle document-related queries efficiently and safely. Always respond as accurately as possible, ensuring your answers comply with ethical guidelines and legal standards. Avoid sharing any information that could be harmful, unethical, biased, or illegal.

#         If a query is ambiguous or unclear, ask clarifying questions to better understand the user's needs before providing a response. If the system cannot retrieve the requested information due to format limitations or content availability, inform the user clearly and suggest possible alternatives.

#         Ensure that your responses enhance user interaction with the system, guiding them on how to upload, retrieve, and manage their documents effectively. If you encounter a question outside your expertise or capability, advise the user accordingly without resorting to misleading or incorrect information.
#         """
        SYSTEM_PROMPT = f"""
       You are a knowledgeable Full stack web developer, responsible for providing code to query . Your responses should be thorough, clear, and aligned with standard object oriented Programming Django python code standards. Always provide Django python code syntax only, while ensuring the code is accurate and relevant to provided project in the document.
Always clarify the context if needed. {query}
        """
        prompt = PromptTemplate(template="""  You are a knowledgeable Full stack web developer, responsible for providing code to query . Your responses should be thorough, clear, and aligned with standard object oriented Programming Django python code standards. Always provide Django python code syntax only, while ensuring the code is accurate and relevant to provided project in the document.
Always clarify the context if needed.

    Context: {context}
    User: {question}
    Full stack web developer:""", input_variables=["context", "question"])

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=db.as_retriever(search_kwargs={"k": 2}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt},
        )

        # Execute a query (you can customize this part as needed)
        query_result = qa_chain(SYSTEM_PROMPT)

                # Execute the query
#         query_result, source_documents = qa_chain.run(query)

        #print("llm!!!!!!!!!!!!!!!!!!!!! Result",query_result)

        # Extract the answer from the 'response' field using string manipulation
        # Assuming the answer always follows "Answer:" and ends at the end of the string
        response_content = query_result['result']
        answer_prefix = "Full stack web developer:"
        answer_start_index = response_content.find(answer_prefix)
        if answer_start_index != -1:
            answer = response_content[answer_start_index + len(answer_prefix):].strip()
            print(answer)
            return answer
        else:
            print("No answer found in the response.")
            return response_content

    def software_developer_agent(self,projectname, file_path, output_path,base_dir):
        sections_queries = {
            #"user classes and characteristics": "Identify the different user classes and their characteristics of the provided project document. Use headings for each user class and bullet points to describe their characteristics.",
#            "Django Directory": """list down django models according to database of project provided in document. """,
#            "Django datamodel": {"query":"Generate Django data model code based on the database models provided in document in python syntax only. Include appropriate field types, relationships (ForeignKey, OneToOneField, ManyToManyField), and any necessary methods.","path":'Knowldgebase/SA/ERD_Diagram.docx'},
            "Django database": {    "query": "Generate SQL database schema code based on the database models and use cases provided in the document. Use SQL syntax only. Include appropriate field types, relationships (ForeignKey, OneToOneField, ManyToManyField), indexes, constraints, and any necessary methods.",
    "path": file_path},
            "Django model": {     "query": "Generate Django model class code based on the database models provided in the document. Use Python syntax only. Include appropriate field types, relationships (ForeignKey, OneToOneField, ManyToManyField), indexes, constraints, and any necessary methods.",
    "path": 'Knowledgebase/SA/SQL_Database.docx'},
            "Django form": {    "query": "Generate Django form class code based on the provided Django model code in the document. Use Python syntax only. Include appropriate form fields, validations, custom methods, and Meta class for model form configuration.",
    "path": 'Knowledgebase/SD/Django_model.docx'},
            "Django serializer": {    "query": "Generate Django serializer class code based on the provided Django model code in the document. Use Python syntax only. Include appropriate serializer fields, validations, custom methods, and Meta class for model serializer configuration.",
    "path": 'Knowledgebase/SD/Django_model.docx'},
            "Django views": {    "query": "Generate Django views code based on the provided Django form code in the document. Use Python syntax only. Include appropriate view classes or functions, decorators, methods for handling requests and responses, and context data for rendering templates.",
    "path": 'Knowledgebase/SD/Django_form.docx'},
            "Django URLs": { "query": "Generate Django URL configuration code based on the provided Django views code in the document. Use Python syntax only. Include appropriate URL patterns, views, names for URL reversing, and any necessary URL parameters.",
    "path": 'Knowledgebase/SD/Django_views.docx'},
            "Django template": {"query":"Generate Django template code (HTML) based on Django views code in HTML syntax only. Include appropriate HTML structure, template tags, and any necessary CSS classes.","path":'Knowldgebase/SD/Django_views.docx'},
#            "CSS": "Generate CSS code based on the based on Django template code (HTML) in CSS syntax only. Include appropriate styles for layout, typography, colors, and any necessary classes or IDs.",
#            "Django template with CSS": "Integrate the following CSS into the Django template. Include the CSS within the appropriate HTML structure and ensure that the template is correctly linked to the CSS styles.""",
    #"unit test code for Django use cases": {"query":"""Generate unit test code for Django use cases based on the following pseudocode. Include tests for each use case scenario, including setup, execution, and verification of expected outcomes.""","path":
        }

        project_info = {}
        project_details=""
        for section, details in sections_queries.items():
            print("query",details.get('query'))
            query = f"For Project {projectname}, {details.get('query')}"
            response = self.generate_response_from_document(details.get("path"), details.get('query'), project_info)
            if section == "Django Directory":
                self.create_directory_structure_from_text(response, base_dir)
            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key]= response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)
    
            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path= f'Knowldgebase/SD/{section_key}.docx'
            # Save the document
            doc.save(path)            
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"


        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project Business analysis Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])
        
        # Save the document
        doc.save(output_path)
        return f"SRS document saved to {output_path}"

    def software_developer_agent_unittests(self,projectname, file_path, output_path,base_dir):
        sections_queries = {
            "unit test code for Django models": """Generate unit test code for Django models based on the following pseudocode. Include tests for model creation, field validation, and any custom methods.""",
            "unit test code for Django forms": """Generate unit test code for Django forms based on the following pseudocode. Include tests for form validation, field validation, and any custom methods.
""",
            "unit test code for Django views": """Generate unit test code for Django views based on the following pseudocode. Include tests for GET and POST request handling, form submissions, and response status codes.
""",
            "unit test code for Django URLs": """Generate unit test code for Django URLs based on the following pseudocode. Include tests for URL resolution, view linking, and response status codes.
""",
            "unit test code for Django templates": """Generate unit test code for Django templates based on the following pseudocode. Include tests for template rendering, context variables, and template tags.
""",
                        "unit test code for Django use cases": """Generate unit test code for Django use cases based on the following pseudocode. Include tests for each use case scenario, including setup, execution, and verification of expected outcomes.
""",
        }

        project_info = {}
        project_details=""
        for section, query in sections_queries.items():
            query = f"For Project {projectname}, {query}"
            response = self.generate_response_from_document(file_path, query, project_info)
            if section == "Django Directory":
                self.generate_project_structure_from_response(response, base_dir)
            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key]= response
            
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"


        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project Uml diagram and Pseudocode Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])
        
        # Save the document
        doc.save(output_path)
        return f"SRS document saved to {output_path}"

    def find_json_structure(self,response):
        try:
            # Attempt to find the JSON structure in the response
            json_start = response.index('{')
            json_end = response.rindex('}') + 1
            json_data = response[json_start:json_end]
            return json_data
        except ValueError:
            raise ValueError("No valid JSON structure found in the response")

    def create_directories(self, base_dir, structure):
        for directory in structure['directories']:
            dir_path = os.path.join(base_dir, directory['name'])
            os.makedirs(dir_path, exist_ok=True)
            
            if 'files' in directory:
                for file_info in directory['files']:
                    if 'contents' in file_info:
                        sub_dir_path = os.path.join(dir_path, file_info['name'])
                        os.makedirs(sub_dir_path, exist_ok=True)
                        for content in file_info['contents']:
                            file_path = os.path.join(sub_dir_path, content)
                            with open(file_path, 'w') as file:
                                pass  # Create an empty file
                    else:
                        file_path = os.path.join(dir_path, file_info['name'])
                        os.makedirs(os.path.dirname(file_path), exist_ok=True)
                        with open(file_path, 'w') as file:
                            pass  # Create an empty file

    def generate_project_structure_from_response(self, response, base_dir):
        json_data = self.find_json_structure(response)
        structure = json.loads(json_data)
        self.create_directories(base_dir, structure)

    def create_directory_structure_from_text(self,response, base_dir):
        lines = response.strip().split('\n')
        current_path = base_dir
        path_stack = []
    
        for line in lines:
            # Ignore empty lines
            if not line.strip():
                continue
    
            # Calculate the level of indentation
            indent_level = len(re.match(r"^\s*\*\s*", line).group()) - 2
    
            # Extract the name and type of the file/directory
            name_type_match = re.match(r"^\s*-\s*(.*) \((.*)\)", line.strip())
            if name_type_match:
                name, file_type = name_type_match.groups()
            else:
                name = line.strip().split(' ')[-1]
                file_type = "directory"
    
            # Calculate the current depth
            while path_stack and indent_level <= path_stack[-1][1]:
                path_stack.pop()
    
            # Get the directory/file name
            current_path = os.path.join(path_stack[-1][0], name) if path_stack else os.path.join(base_dir, name)
    
            # Create directory or file
            if file_type == "directory":
                os.makedirs(current_path, exist_ok=True)
            else:
                os.makedirs(os.path.dirname(current_path), exist_ok=True)
                with open(current_path, 'w') as file:
                    pass  # Create an empty file
    
            # Update the stack
            path_stack.append((current_path, indent_level))