import os
import docx
import re

class CodeSeparator:
    def __init__(self, docx_path, directory):
        self.docx_path = docx_path
        self.directory = directory
        self.text = self.read_docx()
        self.html_code = ""
        self.css_code = ""
        self.js_code = ""

    def read_docx(self):
        doc = docx.Document(self.docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)

    def separate_code(self):
        # Regular expressions to find HTML, CSS, and JS code blocks
        html_pattern = re.compile(r'<!DOCTYPE html>.*<\/html>', re.DOTALL | re.IGNORECASE)
        css_pattern = re.compile(r'\/\*.*?Global Styles.*?\*\/.*', re.DOTALL)
        js_pattern = js_pattern = re.compile(r'(<\s*script[^>]*>.*?<\s*\/\s*script\s*>|\/\/.*$)', re.DOTALL | re.IGNORECASE | re.MULTILINE)

        html_match = html_pattern.search(self.text)
        css_match = css_pattern.search(self.text)
        js_match = js_pattern.search(self.text)

        if html_match:
            self.html_code = html_match.group().strip()
        if css_match:
            self.css_code = css_match.group().strip()
        if js_match:
            self.js_code = js_match.group().strip()

    def save_to_files(self):
        os.makedirs(self.directory, exist_ok=True)

        html_path = os.path.join(self.directory, "index.html")
        css_path = os.path.join(self.directory, "style.css")
        js_path = os.path.join(self.directory, "script.js")

        if self.html_code:
            with open(html_path, 'w') as file:
                file.write(self.html_code)
        
        if self.css_code:
            with open(css_path, 'w') as file:
                file.write(self.css_code)
        
        if self.js_code:
            with open(js_path, 'w') as file:
                file.write(self.js_code)

import re
import os
from docx import Document

def extract_and_create_html_files(docx_path, project_directory):
    """
    Extracts HTML content from a .docx file and creates individual HTML files in the specified project directory.
    
    Args:
        docx_path (str): The path to the .docx file containing the combined HTML content.
        project_directory (str): The directory where the extracted HTML files will be stored.
    """
    # Ensure the project directory exists
    if not os.path.exists(project_directory):
        os.makedirs(project_directory)

    # Read the combined HTML content from the .docx file
    document = Document(docx_path)
    combined_html_content = "\n".join(paragraph.text for paragraph in document.paragraphs)

    # Regular expression pattern to match each HTML file's content
    pattern = re.compile(r'\*\*(.+?\.html)\*\*\s*```html(.*?)```', re.DOTALL)

    # Find all matches
    matches = pattern.findall(combined_html_content)

    # Create each HTML file with the extracted content
    for filename, content in matches:
        file_path = os.path.join(project_directory, filename)
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content.strip())
        print(f"Created file: {file_path}")


def extract_and_create_css_file(docx_path, project_directory):
    """
    Extracts CSS content from a .docx file and creates a styles.css file in the specified project directory.
    
    Args:
        docx_path (str): The path to the .docx file containing the CSS content.
        project_directory (str): The directory where the styles.css file will be stored.
    """
    # Ensure the project directory exists
    if not os.path.exists(project_directory):
        os.makedirs(project_directory)

    # Read the content from the .docx file
    document = Document(docx_path)
    combined_content = "\n".join(paragraph.text for paragraph in document.paragraphs)

    # Regular expression pattern to detect CSS content
    css_pattern = re.compile(r'([^{]+\{[^}]+\})', re.DOTALL)

    # Find all matches
    matches = css_pattern.findall(combined_content)

    # Combine the matched CSS content
    css_content = "\n".join(matches)

    # Create the styles.css file with the extracted content
    css_file_path = os.path.join(project_directory, 'styles.css')
    with open(css_file_path, 'w', encoding='utf-8') as file:
        file.write(css_content)
    print(f"Created file: {css_file_path}")