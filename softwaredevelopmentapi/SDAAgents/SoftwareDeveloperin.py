import re
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, TextStreamer, pipeline
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain_community.document_loaders import Docx2txtLoader
from langchain.schema import Document
from docx import Document as DocxDocument
import os
from softwaredevelopmentapi.model_loader import llm_model,tokenizer
class DjangoModelCodeGenerator:
    def __init__(self, doc_path, model_id, embedding_model, output_dir='output', device='cuda'):
        self.tokenizer = tokenizer
        self.llm = llm_model
        
        self.doc_path = doc_path
        self.model_id = model_id
        self.embedding_model = embedding_model
        self.device = device
        self.output_dir = output_dir
        self.data_string = None
        self.model_names = None
        self.texts = None
        self.db = None
        self.qa_chain = None

        # Ensure the output directory exists
        os.makedirs(self.output_dir, exist_ok=True)

    def load_document(self):
        loader = Docx2txtLoader(self.doc_path)
        data = loader.load()
        self.data_string = ''.join([doc.page_content for doc in data]) if isinstance(data, list) else data.page_content
        print("Extracted Text:")
        print(self.data_string)

    def extract_model_names(self):
        pattern = r'CREATE TABLE (\w+)'
        self.model_names = re.findall(pattern, self.data_string)
        print("Extracted Model Names:")
        print(self.model_names)

    def set_up_embeddings(self):
        embeddings = HuggingFaceEmbeddings(
            model_name=self.embedding_model,
            model_kwargs={"device": self.device}
        )
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
        document = Document(page_content=self.data_string, metadata={})
        self.texts = text_splitter.split_documents([document])
        print("Number of text chunks:", len(self.texts))

        self.db = Chroma.from_documents(self.texts, embeddings)

    def set_up_language_model(self):
        tokenizer = self.tokenizer
        model = self.llm
        # tokenizer = AutoTokenizer.from_pretrained(self.model_id, use_auth_token="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")
        # model = AutoModelForCausalLM.from_pretrained(
        #     self.model_id,
        #     torch_dtype=torch.bfloat16,
        #     device_map="auto",
        #     use_auth_token="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",
        # )
        streamer = TextStreamer(tokenizer, skip_prompt=True, skip_special_tokens=True)
        
        text_pipeline = pipeline(
            "text-generation",
            model=model,
            tokenizer=tokenizer,
            max_new_tokens=1024,
            temperature=0.5,
            top_p=0.95,
            repetition_penalty=1.15,
            streamer=streamer,
        )

        self.llm = HuggingFacePipeline(pipeline=text_pipeline, model_kwargs={"temperature": 0.7})

    def create_qa_chain(self):
        prompt_template = PromptTemplate(template="""You are a highly knowledgeable and professional Fullstack web_developer chatbot, dedicated to providing accurate and detailed informative answers to user queries. Please don't generate irrelevant or extra information which the user did not ask you to generate.    
Context: {context}
User: {question}
Chatbot:""", input_variables=["context", "question"])

        self.qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=self.db.as_retriever(search_kwargs={"k": 2}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt_template},
        )

    def generate_code(self):
        prompts = []
        answers = []
        for model_name in self.model_names:
            prompt = f"Write Django (model.py,form.py,serilizer.py,views.py,) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for."
            result = self.qa_chain(prompt)
            prompts.append(prompt)
            answers.append(result)
        
        return prompts, answers

    def save_to_docx(self, model_name, content,output_path):
        try:
            doc = DocxDocument()
            doc.add_heading(f'Code for {model_name}', level=1)
            doc.add_paragraph(content)
            # file_path = os.path.join(self.output_dir, f'{model_name}.docx')
            doc.save(output_path)
            print(f'Saved {output_path}')
            # return 'Code has been generated.'
        except Exception as e:
            print(f"Failed to save {model_name}.docx: {e}")
            # return "Failed to generated code."


    def display_and_save_results(self, prompts, answers):
        try:
            for i, model_name in enumerate(self.model_names):
                print(f"Prompt for {model_name}:")
                print(prompts[i])
                print(f"Answer for {model_name}:")
                print(answers[i])
                print("\n")
                self.save_to_docx(model_name, answers[i]['result'])
            return 'Code has been generated.'
        except Exception as e:
            print(f"An error occurred: {e}")
            return 'Failed to generate code.'


# def software_developer_agent(input_path,output_path,projectid):
    
#     generator = DjangoModelCodeGenerator(
#     doc_path=input_path,
#     model_id="meta-llama/Meta-Llama-3-8B-Instruct",
#     embedding_model="sentence-transformers/all-MiniLM-L6-v2"
#     )
#     generator.load_document()
#     generator.extract_model_names()
#     generator.set_up_embeddings()
#     generator.set_up_language_model()
#     generator.create_qa_chain()
#     prompts, answers = generator.generate_code()
#     generator.display_and_save_results(prompts, answers)
        





# Usage
# generator = DjangoModelCodeGenerator(
#     doc_path="knowledgebase/SD/Django_datamodel1.docx",
#     model_id="meta-llama/Meta-Llama-3-8B-Instruct",
#     embedding_model="sentence-transformers/all-MiniLM-L6-v2"
# )
# generator.load_document()
# generator.extract_model_names()
# generator.set_up_embeddings()
# generator.set_up_language_model()
# generator.create_qa_chain()
# prompts, answers = generator.generate_code()
# generator.display_and_save_results(prompts, answers)