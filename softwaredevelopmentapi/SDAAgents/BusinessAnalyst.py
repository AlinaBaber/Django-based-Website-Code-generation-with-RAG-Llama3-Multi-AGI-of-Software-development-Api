from huggingface_hub import login
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import json
import textwrap
import os
#####################
import os
import shutil
import torch
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline
import torch
from docx import Document
from softwaredevelopmentapi.model_loader import llm_model,tokenizer

class BusinessAnalystAgent:
    def __init__(self, model_name="meta-llama/Llama-2-7b-chat-hf", tokenizer_name="meta-llama/Llama-2-7b-chat-hf"):
        #self.tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True,
         #                                              use_auth_token="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")
        #self.model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,
        #                                                  use_auth_token="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")
        self.tokenizer = tokenizer
        self.model = llm_model
        self.db = None

    def read_pdf(self, file_path):
        loader = PyPDFLoader(file_path)
        data = loader.load()

        return data

    def read_docx(self, file_path):
        loader = Docx2txtLoader(file_path)
        data = loader.load()

        return data

    def read_txt(self, file_path):
        loader = TextLoader(file_path)
        data = loader.load()

        return data

    def add_text_to_docx(self, file_path, text, output_path):
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(output_path)
        
    def load_db(self, file_path):
        # if len(project_info) > 0:
        #     self.add_text_to_docx(file_path, project_info, "Knowldgebase/temp/temp.docx")
        """Generate a response based on a user query from a document."""
        # Determine the file type and extract text
        if file_path.lower().endswith('.pdf'):
            document_text = self.read_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            document_text = self.read_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            document_text = self.read_txt(file_path)
        else:
            raise ValueError("Unsupported file type. Supported types are: 'pdf', 'docx', 'txt'.")

        self.streamer = TextStreamer(self.tokenizer, skip_prompt=True, skip_special_tokens=True)

        self.text_pipeline = pipeline(
            "text-generation",
            model=self.model,
            tokenizer=self.tokenizer,
            max_new_tokens=1024,
            temperature=0.5,
            top_p=0.95,
            repetition_penalty=1.15,
            streamer=self.streamer,
        )

        # document_text[0]= f"\n{document_text[0]=}\n{project_info}"

        self.llm = HuggingFacePipeline(pipeline=self.text_pipeline, model_kwargs={"temperature": 0.5})

        # Step 2: Split the text into manageable chunks
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
        self.texts = self.text_splitter.split_documents(document_text)

        # Step 3: Generate embeddings for the texts
        self.embeddings = HuggingFaceEmbeddings(
            model_name="hkunlp/instructor-xl", model_kwargs={"device": "cuda"}
        )
        self.db = Chroma.from_documents(self.texts, self.embeddings)
    def generate_response_from_document(self,query):
        SYSTEM_PROMPT = f"""
       You are a knowledgeable chatbot, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with business analysis standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant.
Always clarify the context if needed. the query is {query}
        """
        prompt = PromptTemplate(template=""" You are a knowledgeable chatbot, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with business analysis standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant. Answer should be in English
Always clarify the context if needed. 

    Context: {context}
    User: {question}
    chatbot:""", input_variables=["context", "question"])

        qa_chain = RetrievalQA.from_chain_type(
            
            llm=self.llm,
            chain_type="stuff",
            retriever=self.db.as_retriever(search_kwargs={"k": 2}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt},
        )

        # Execute a query (you can customize this part as needed)
        query_result = qa_chain(query)

        # Execute the query
        #         query_result, source_documents = qa_chain.run(query)

        # print("llm!!!!!!!!!!!!!!!!!!!!! Result",query_result)

        # Extract the answer from the 'response' field using string manipulation
        # Assuming the answer always follows "Answer:" and ends at the end of the string
        response_content = query_result['result']
        answer_prefix = "Chatbot: "
        answer_start_index = response_content.find(answer_prefix)
        if answer_start_index != -1:
            answer = response_content[answer_start_index + len(answer_prefix):].strip()
            print(answer)
            return answer
        else:
            print("No answer found in the response.")
            return response_content        
#     def generate_response_from_document(self, file_path, query, project_info):
#         #if len(project_info) > 0:
#         #    self.add_text_to_docx(file_path, project_info, "Knowldgebase/temp/temp.docx")
#         """Generate a response based on a user query from a document."""
#         # Determine the file type and extract text
#         if file_path.lower().endswith('.pdf'):
#             document_text = self.read_pdf(file_path)
#         elif file_path.lower().endswith('.docx'):
#             document_text = self.read_docx(file_path)
#         elif file_path.lower().endswith('.txt'):
#             document_text = self.read_txt(file_path)
#         else:
#             raise ValueError("Unsupported file type. Supported types are: 'pdf', 'docx', 'txt'.")

#         streamer = TextStreamer(self.tokenizer, skip_prompt=True, skip_special_tokens=True)

#         text_pipeline = pipeline(
#             "text-generation",
#             model=self.model,
#             tokenizer=self.tokenizer,
#             max_new_tokens=1024,
#             temperature=0.5,
#             top_p=0.95,
#             repetition_penalty=1.15,
#             streamer=streamer,
#         )

#         # document_text[0]= f"\n{document_text[0]=}\n{project_info}"

#         llm = HuggingFacePipeline(pipeline=text_pipeline, model_kwargs={"temperature": 0.5})

#         # Step 2: Split the text into manageable chunks
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
#         texts = text_splitter.split_documents(document_text)

#         # Step 3: Generate embeddings for the texts
#         embeddings = HuggingFaceEmbeddings(
#             model_name="hkunlp/instructor-xl", model_kwargs={"device": "cuda"}
#         )
#         db = Chroma.from_documents(texts, embeddings)

#         #         SYSTEM_PROMPT = """
#         #         You are a knowledgeable and precise assistant designed to handle document-related queries efficiently and safely. Always respond as accurately as possible, ensuring your answers comply with ethical guidelines and legal standards. Avoid sharing any information that could be harmful, unethical, biased, or illegal.

#         #         If a query is ambiguous or unclear, ask clarifying questions to better understand the user's needs before providing a response. If the system cannot retrieve the requested information due to format limitations or content availability, inform the user clearly and suggest possible alternatives.

#         #         Ensure that your responses enhance user interaction with the system, guiding them on how to upload, retrieve, and manage their documents effectively. If you encounter a question outside your expertise or capability, advise the user accordingly without resorting to misleading or incorrect information.
#         #         """
#         SYSTEM_PROMPT = f"""
#        You are a knowledgeable chatbot, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with business analysis standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant.
# Always clarify the context if needed. the query is {query}
#         """
#         prompt = PromptTemplate(template=""" You are a knowledgeable chatbot, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with business analysis standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant.
# Always clarify the context if needed. 

#     Context: {context}
#     User: {question}
#     chatbot:""", input_variables=["context", "question"])

#         qa_chain = RetrievalQA.from_chain_type(
#             llm=llm,
#             chain_type="stuff",
#             retriever=db.as_retriever(search_kwargs={"k": 2}),
#             return_source_documents=True,
#             chain_type_kwargs={"prompt": prompt},
#         )

#         # Execute a query (you can customize this part as needed)
#         query_result = qa_chain(query)

#         # Execute the query
#         #         query_result, source_documents = qa_chain.run(query)

#         # print("llm!!!!!!!!!!!!!!!!!!!!! Result",query_result)

#         # Extract the answer from the 'response' field using string manipulation
#         # Assuming the answer always follows "Answer:" and ends at the end of the string
#         response_content = query_result['result']
#         answer_prefix = "chatbot: "
#         answer_start_index = response_content.find(answer_prefix)
#         if answer_start_index != -1:
#             answer = response_content[answer_start_index + len(answer_prefix):].strip()
#             print(answer)
#             return answer
#         else:
#             print("No answer found in the response.")
#             return response_content

    def business_analyst_agent(self,file_path, output_path,projectid):
        sections_queries = {
            "title": "What is the project title as mentioned in the provided project document?",
            "description": "Provide a detailed description of the project from the provided document, including its background, purpose, and scope. Use paragraphs for in-depth explanations and bullet points for key highlights.",
            "project objectives and goals": "What are the project goals and objectives based on the provided project document? Provide a comprehensive list with detailed explanations, organized with bullet points.",
            "tasks": "List the tasks involved in the project as per the provided document. Organize the tasks in bullet points and provide brief descriptions for each.",
            #       project objectives": "Detail the project goals and objectives from the provided document. Use headings for different goal categories and bullet points for specific objectives.",
            "business case": "Provide the business case for the project as outlined in the document, including the rationale, benefits, and potential impacts. Use paragraphs for detailed explanations and bullet points for key points.",
            "functional requirements": "List the functional requirements for the project as mentioned in the document. Organize the requirements with headings for different functional areas and bullet points for specific requirements.",
            "non-functional requirements": "List the non-functional requirements for the project from the document, such as performance, usability, and reliability. Use headings for different categories and bullet points for specific requirements.",
            "user stories & acceptance criteria": "List the user stories and acceptance criteria based on the project document. Use bullet points to list each user story and include brief descriptions.",
#            "stakeholders information": "Provide detailed information about the stakeholders, including their roles and interests as mentioned in the project document. Use headings for different stakeholder groups and bullet points for specific details.",
            #        "current systems overview": "Give an overview of the current systems, including their functions and limitations as described in the project document. Use paragraphs for detailed explanations and bullet points for key points.",
#            "technical constraints": "List any technical constraints for the project as mentioned in the document. Use bullet points to clearly list each constraint and provide brief explanations.",
            #        "integration points": "Identify the integration points in the project as outlined in the document. Use bullet points to list each integration point and provide brief descriptions.",
            "in-scope features": "What are the in-scope features for the project as defined in the document? Use bullet points to list each feature and provide brief explanations.",
            #        "out-of-scope features": "What are the out-of-scope features for the project as defined in the document? Use bullet points to list each feature and provide brief explanations.",
            "data models": "Provide the detailed data models relevant to the project as mentioned in the document. Use headings for different data model types and bullet points for specific details.",
            # "data sources": "What are the data sources for the project as described in the document? Use bullet points to list each data source and provide brief descriptions.",
            # "data security": "Detail the data security measures for the project as outlined in the document. Use paragraphs for detailed explanations and bullet points for key measures.",
#            "user profiles": "Provide user profiles relevant to the project as mentioned in the document. Use headings for different user types and bullet points for specific characteristics.",
            # "user experience requirements": "What are the user experience requirements for the project as outlined in the document? Use headings for different UX aspects and bullet points for specific requirements.",
            # "legal requirements": "List the legal requirements for the project as mentioned in the document. Use bullet points to clearly list each requirement and provide brief explanations.",
            # "compliance standards": "Detail the compliance standards applicable to the project as described in the document. Use headings for different standards and bullet points for specific requirements.",
            # "risks and mitigations": "Identify the risks and their mitigations for the project as outlined in the document. Use bullet points to list each risk and its corresponding mitigation strategy.",
            #  "assumptions and dependencies": "List the assumptions and dependencies for the project as mentioned in the document. Use bullet points to clearly list each assumption and dependency and provide brief explanations.",
#            "deployment requirements": "Provide the deployment requirements for the project as outlined in the document. Use paragraphs for detailed explanations and bullet points for key requirements.",
            #        "maintenance and support": "What are the maintenance and support plans for the project as outlined in the document? Use paragraphs for detailed explanations and bullet points for key plans.",
            #        "project timeline": "Outline the project timeline as described in the document. Use headings for different phases and bullet points for specific milestones.",
            #        "budget": "Provide the project budget details as mentioned in the document. Use paragraphs for detailed explanations and bullet points for key budget items.",
            #        "documentation": "List the documentation requirements for the project as outlined in the document. Use bullet points to list each required document and provide brief descriptions.",
            #        "change management process": "Describe the change management process for the project as mentioned in the document. Use paragraphs for detailed explanations and bullet points for key steps.",
            #        "communication channels": "What are the communication channels for the project as described in the document? Use bullet points to list each channel and provide brief descriptions.",
            #        "reporting requirements": "Detail the reporting requirements for the project as mentioned in the document. Use paragraphs for detailed explanations and bullet points for key reports."
        }
        self.load_db(file_path)
        project_info = {}
        project_details = ""
        for section, query in sections_queries.items():
            #query = f"For {projectname}, {query}"
            response = self.generate_response_from_document(query)

            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key] = response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)

            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path = f'Knowldgebase/BA/{section_key}{projectid}.docx'
            # Save the document
            doc.save(path)
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"
                        # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()

        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project Business analysis Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"BA document saved to {output_path}"
        return "Business Analysis has been completed"