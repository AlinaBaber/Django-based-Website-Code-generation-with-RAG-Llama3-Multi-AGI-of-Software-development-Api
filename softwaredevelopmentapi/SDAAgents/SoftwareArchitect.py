from huggingface_hub import login
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import json
import textwrap
import os
#####################
import os
import shutil
import torch
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline
import torch
from docx import Document
from softwaredevelopmentapi.model_loader import llm_model,tokenizer

class SoftwareArchitectAgent:
    def __init__(self, model_name="facebook/llama-2", tokenizer_name="facebook/llama-2"):
        #self.tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True,
        #                                               use_auth_token="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")
        #self.model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,
        #                                                  use_auth_token="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")
        self.tokenizer = tokenizer
        self.model = llm_model
        self.db = None

    def read_pdf(self, file_path):
        loader = PyPDFLoader(file_path)
        data = loader.load()

        return data

    def read_docx(self, file_path):
        loader = Docx2txtLoader(file_path)
        data = loader.load()

        return data

    def read_txt(self, file_path):
        loader = TextLoader(file_path)
        data = loader.load()

        return data

    def add_text_to_docx(self, file_path, text, output_path):
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(output_path)
    def load_db(self, file_path):
        # if len(project_info) > 0:
        #     self.add_text_to_docx(file_path, project_info, "Knowldgebase/temp/temp.docx")
        """Generate a response based on a user query from a document."""
        # Determine the file type and extract text
        if file_path.lower().endswith('.pdf'):
            document_text = self.read_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            document_text = self.read_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            document_text = self.read_txt(file_path)
        else:
            raise ValueError("Unsupported file type. Supported types are: 'pdf', 'docx', 'txt'.")

        self.streamer = TextStreamer(self.tokenizer, skip_prompt=True, skip_special_tokens=True)

        self.text_pipeline = pipeline(
            "text-generation",
            model=self.model,
            tokenizer=self.tokenizer,
            max_new_tokens=1024,
            temperature=0.5,
            top_p=0.95,
            repetition_penalty=1.15,
            streamer=self.streamer,
        )

        # document_text[0]= f"\n{document_text[0]=}\n{project_info}"

        self.llm = HuggingFacePipeline(pipeline=self.text_pipeline, model_kwargs={"temperature": 0.5})

        # Step 2: Split the text into manageable chunks
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
        self.texts = self.text_splitter.split_documents(document_text)

        # Step 3: Generate embeddings for the texts
        self.embeddings = HuggingFaceEmbeddings(
            model_name="hkunlp/instructor-xl", model_kwargs={"device": "cuda"}
        )
        self.db = Chroma.from_documents(self.texts, self.embeddings)
        
        return self.streamer,self.text_pipeline,self.text_splitter,self.texts,self.embeddings,self.db
    def generate_response_from_document(self,query):
        SYSTEM_PROMPT = f"""
       You are a knowledgeable Software Developer Assistant, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with software engineering standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant to provided project in the document.
Always clarify the context if needed. the query is {query}
        """
        prompt = PromptTemplate(template=""" You are a knowledgeable AI Assistant, responsible for providing detailed and comprehensive answer to query . Your responses should be thorough, clear, and aligned with software engineering standards. Always provide as much detail as possible, while ensuring the information is accurate and relevant to provided project in the document.
Always clarify the context if needed.

    Context: {context}
    User: {question}
    Chatbot:""", input_variables=["context", "question"])

        qa_chain = RetrievalQA.from_chain_type(
            
            llm=self.llm,
            chain_type="stuff",
            retriever=self.db.as_retriever(search_kwargs={"k": 2}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt},
        )

        # Execute a query (you can customize this part as needed)
        query_result = qa_chain(query)

        # Execute the query
        #         query_result, source_documents = qa_chain.run(query)

        # print("llm!!!!!!!!!!!!!!!!!!!!! Result",query_result)

        # Extract the answer from the 'response' field using string manipulation
        # Assuming the answer always follows "Answer:" and ends at the end of the string
        response_content = query_result['result']
        answer_prefix = "Chatbot: "
        answer_start_index = response_content.find(answer_prefix)
        if answer_start_index != -1:
            answer = response_content[answer_start_index + len(answer_prefix):].strip()
            print(answer)
            return answer
        else:
            print("No answer found in the response.")
            return response_content

    def software_architect_agent(self, file_path, output_path,projectid):
        sections_queries = {
#            "title": "What is the project title as mentioned in the provided project document?",
            "Project Summery": "List all a detailed SUmmery of the project from the provided document, including its background, purpose, and scope. Use paragraphs for in-depth explanations and bullet points for key highlights.",
#            "project objectives and goals": "What are the project goals and objectives based on the provided project document? Provide a comprehensive list with detailed explanations, organized with bullet points.",
#            "tasks": "List the tasks involved in the project as per the provided document. Organize the tasks in bullet points and provide brief descriptions for each.",
            "user stories & acceptance criteria": "List the user stories and acceptance criteria based on the project document. Use bullet points to list each user story and include brief descriptions.",
            # "user classes and characteristics": "Identify the different user classes and their characteristics of the provided project document. Use headings for each user class and bullet points to describe their characteristics.",
#            "operating environment": "Detail the operating environment where the provided project in the document will be used. Use paragraphs for detailed descriptions and bullet points for key environmental factors.",
            #            "design and implementation constraints": "Detail the design and implementation constraints of the provided project document. Use bullet points to clearly list each constraint and provide brief explanations.",
            # "assumptions and dependencies": "List any assumptions and dependencies that impact of the provided project document. Use bullet points for each assumption and dependency and provide brief explanations.",
            "functional requirements": " List all the functional requirements of the system of the provided project document. Use headings for different functional areas and bullet points to list specific requirements.",
            #            "external interface requirements": "Detail the all requirements for external interfaces of the provided project document. Use headings for different types of interfaces and bullet points for specific requirements.",
            "user interfaces": "List all the requirements for user interfaces UI/UX of the provided project document. Use headings for different interface types and bullet points for specific requirements.",
#            "system features": "Detail the system features, including the main functionalities of the provided project document. Use headings for different feature categories and bullet points for individual features.",
            "use case": "List all Detailed usecases of the provided project document. its should have all the website pages and button usecases as well .Use headings for different interface types and bullet points for specific requirements.",
            #            "hardware interfaces": "Describe the hardware interface requirements of the provided project document. Use headings for different hardware components and bullet points for specific requirements.",
            #            "software interfaces": "Detail the software interface requirements of the provided project document. Use headings for different software components and bullet points for specific requirements.",
            #            "communication interfaces": "Describe the communication interface requirements of the provided project document. Use headings for different communication protocols and bullet points for specific requirements.",
             "non-functional requirements": "List all non-functional requirements, including performance, usability, reliability,safety andsecurity of the provided project document. Use headings for different non-functional categories and bullet points for specific requirements.",
            "business rules": "Detail the business rules that apply to the project of the provided project document. Use headings for different business rules categories and bullet points for specific rules.",
            "Database": """Generate detailed big database tables of the project according to information provided in document in SQL syntax. Include all major entities, their attributes, primary keys, and relationships (one-to-one, one-to-many, many-to-many) in SQL syntax. tables should be more than 20 """,
#            "HTMLPages": """Based on the provided SRS document which includes detailed business rules, functional requirements, user interface specifications, and system features, please provide a detailed list of each HTML pages that are required for the project in to json format. For each page, include a brief description of its purpose, key features, and any specific elements or sections that should be included.""",
        }
        self.load_db(file_path)
        project_info = {}
        project_details = ""
        for section, query in sections_queries.items():
            #query = f"For Project {projectname}, {query}"
            response = self.generate_response_from_document(query)

            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key] = response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)

            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path = f'Knowldgebase/SA/{section_key}{projectid}.docx'
            # Save the document
            doc.save(path)
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"
                        # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()

        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project SRS Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"SRS document saved to {output_path}"
        return "SRS document has been completed."

    def software_architect_agent_design(self, file_path, output_path,projectid):
        sections_queries = {
            "HTMLPages": """Based on the provided SRS document which includes detailed business rules, functional requirements, user interface specifications, and system features, please provide a detailed list of each HTML pages that are required for the project in to json format. For each page, include a brief description of its purpose, key features, and any specific elements or sections that should be included.""",
#            "flow": """Generate a low-level process flow for the project according to information provided in document into json format. Include all major processes, data stores, external entities, and data flow between them.""",
#                        "flow diagram": """Generate a low-level process flow diagram for the project according to information provided in document in Mermaid syntax . Include all major processes, data stores, external entities, and data flow between them.""",
            "usecase": """Generate a low-level usecases of the project according to use cases provided in documentinto json format. each usecase should include  name , description , actor, association.""",
            #            "usecase diagram": """Generate a low-level use case diagram of the project according to use cases provided in document in Mermaid syntax. Include all primary actors, use cases, and their relationships (associations, includes, extends).""",
            # "sequence": """Generate a low-level sequence of the project for the main interactions in the project provided in document into json format. Include the key participants (objects or classes), messages exchanged, and the sequence of events.""",
            #            "sequence diagram": """Generate a low-level sequence diagram of the project in Mermaid syntax for the main interactions in the project provided in document.. Include the key participants (objects or classes), messages exchanged, and the sequence of events.""",
#            "dataflow": """Generate a data flow  of the project provided in document into json format. Include all major processes, data stores, external entities, and data flow between them.""",
            #            "dataflow diagram": """Generate a data flow diagram of the project in Mermaid syntax. Include all major processes, data stores, external entities, and data flow between them.""",
            "SQL Database":"""Generate detailed big database tables of the project according to information provided in document in SQL syntax. Include all major entities, their attributes, primary keys, and relationships (one-to-one, one-to-many, many-to-many) in SQL syntax. tables should be more than 20. dont miss any table. """,
            #            "ERD Diagram": """Generate an detailed database ERD Diagram of the project according to information provided in document in Mermaid syntax. Include all major entities, their attributes, primary keys, and relationships (one-to-one, one-to-many, many-to-many) in SQL syntax.""",
            # "class ": """Generate a classes  related to each usecase provided in document for the project in the provided in document  into json format. Ensure to include all major classes, their attributes, methods, and relationships (inheritance, association, aggregation, composition).""",
            #            "class diagram": """Generate a class diagram  related to each usecase provided in document for the project in Mermaid syntax. Ensure to include all major classes, their attributes, methods, and relationships (inheritance, association, aggregation, composition).""",
#             "Detailed Pseudocode": """Generate low-level pseudocode in the project in the provided in document into json format. The pseudocode should be clear, step-by-step, and cover all major functionalities and logic. Use a structured format for clarity.
# """,
            # "Deployment": """generate a deployment details of the project in the provided in document according to web platform into json format. Include all hardware nodes, software components, and the relationships between them.""",
            # #            "Deployment Diagram": """generate a deployment diagram of the project according to web platform in Mermaid syntax. Include all hardware nodes, software components, and the relationships between them.
            # """,
        }
        self.load_db(file_path)
        project_info = {}
        project_details = ""
        for section, query in sections_queries.items():
            #query = f"For Project {projectname}, {query}"
            response = self.generate_response_from_document(query)

            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key] = response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)

            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path = f'Knowldgebase/SA/{section_key}{projectid}.docx'
            # Save the document
            doc.save(path)
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"
                        # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()

        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project Uml diagram and Pseudocode Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"SRS document saved to {output_path}"
        return "Design document has been completed."
