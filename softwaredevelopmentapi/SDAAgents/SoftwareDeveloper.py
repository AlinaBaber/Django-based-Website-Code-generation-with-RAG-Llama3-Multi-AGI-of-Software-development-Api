from huggingface_hub import login
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import json
import textwrap
import os
#####################
import os
import shutil
import torch
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain import HuggingFacePipeline, PromptTemplate
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from transformers import AutoModelForCausalLM, AutoTokenizer, TextStreamer, pipeline
import torch
from docx import Document
import os
import json
import os
import re
#from softwaredevelopmentapi.model_loader import llm_model,tokenizer

class SoftwareDeveloperAgent:
    def __init__(self, model_name="meta-llama/Meta-Llama-3-8B-Instruct", tokenizer_name="meta-llama/Meta-Llama-3-8B-Instruct"):
        #self.tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True,
        #                                               use_auth_token="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")
        #self.model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,
        #                                                  use_auth_token="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")
        model_name="meta-llama/Meta-Llama-3-8B-Instruct"
        #model_name="meta-llama/Llama-2-7b-chat-hf"
        llm_model = AutoModelForCausalLM.from_pretrained(model_name, device_map='auto', torch_dtype=torch.float16,use_auth_token="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")
        tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True,
                                  use_auth_token="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")
        self.tokenizer = tokenizer
        self.model = llm_model
        
    def read_pdf(self, file_path):
        loader = PyPDFLoader(file_path)
        data = loader.load()

        return data

    def read_docx(self, file_path):
        loader = Docx2txtLoader(file_path)
        data = loader.load()

        return data

    def read_txt(self, file_path):
        loader = TextLoader(file_path)
        data = loader.load()

        return data

    def add_text_to_docx(self, file_path, text, output_path):
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(output_path)

    def generate_response_from_document(self, file_path, query, project_info):
        #if len(project_info) > 0:
        #    self.add_text_to_docx(file_path, project_info, "Knowldgebase/temp/temp.docx")
        # """Generate a response based on a user query from a document."""
        # Determine the file type and extract text
        if file_path.lower().endswith('.pdf'):
            document_text = self.read_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            document_text = self.read_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            document_text = self.read_txt(file_path)
        else:
            raise ValueError("Unsupported file type. Supported types are: 'pdf', 'docx', 'txt'.")

        streamer = TextStreamer(self.tokenizer, skip_prompt=True, skip_special_tokens=True)

        text_pipeline = pipeline(
            "text-generation",
            model=self.model,
            tokenizer=self.tokenizer,
            max_new_tokens=2000,
            temperature=0.5,
            top_p=0.95,
            repetition_penalty=1.15,
            streamer=streamer,
        )

        # document_text[0]= f"\n{document_text[0]=}\n{project_info}"

        llm = HuggingFacePipeline(pipeline=text_pipeline, model_kwargs={"temperature": 0.7})

        # Step 2: Split the text into manageable chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=64)
        texts = text_splitter.split_documents(document_text)

        # Step 3: Generate embeddings for the texts
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={"device": "cuda"}
        )
        db = Chroma.from_documents(texts, embeddings)

        #         SYSTEM_PROMPT = """
        #         You are a knowledgeable and precise assistant designed to handle document-related queries efficiently and safely. Always respond as accurately as possible, ensuring your answers comply with ethical guidelines and legal standards. Avoid sharing any information that could be harmful, unethical, biased, or illegal.

        #         If a query is ambiguous or unclear, ask clarifying questions to better understand the user's needs before providing a response. If the system cannot retrieve the requested information due to format limitations or content availability, inform the user clearly and suggest possible alternatives.

        #         Ensure that your responses enhance user interaction with the system, guiding them on how to upload, retrieve, and manage their documents effectively. If you encounter a question outside your expertise or capability, advise the user accordingly without resorting to misleading or incorrect information.
        #         """
        SYSTEM_PROMPT = f"""
       You are a knowledgeable AI Assistant, responsible for providing code to query . Your responses should be thorough, clear, and aligned with standard object oriented Programming Django python code standards. Always provide Django python code syntax only, while ensuring the code is accurate and relevant to provided project in the document.
Always clarify the context if needed. {query}
        """
        prompt = PromptTemplate(template="""You are a knowledgeable AI Assistant, responsible for providing code to query . Your responses should be thorough, clear, and aligned with standard object oriented Programming Django python code standards, while ensuring the code is accurate according provided project in the document.   
Context: {context}
User: {question}
AI Assistant:""", input_variables=["context", "question"])

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=db.as_retriever(search_kwargs={"k": 2}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt},
        )

        # Execute a query (you can customize this part as needed)
        query_result = qa_chain(query)

        # Execute the query
        #         query_result, source_documents = qa_chain.run(query)

        # print("llm!!!!!!!!!!!!!!!!!!!!! Result",query_result)

        # Extract the answer from the 'response' field using string manipulation
        # Assuming the answer always follows "Answer:" and ends at the end of the string
        response_content = query_result['result']
        answer_prefix = "AI Assistant:"
        answer_start_index = response_content.find(answer_prefix)
        if answer_start_index != -1:
            answer = response_content[answer_start_index + len(answer_prefix):].strip()
            print(answer)
            return answer
        else:
            print("No answer found in the response.")
            return response_content

    def software_developer_agent_permodel(self, file_path, output_path, base_dir,projectid,model_name):
        sections_queries = {
            "Django model": {
                "query": f"Write Django (model.py) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": file_path,
                "outpath":  f'{base_dir}/{model_name}.docx'},
            "Django form": {
                "query": f"Write Django (form.py) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{model_name}.docx',
                "outpath":  f'{base_dir}/{model_name}form.docx'},
            "Django serializer": {
                "query": f"Write Django (serilizer.py) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{model_name}.docx',
                "outpath":  f'{base_dir}/{model_name}Serializer.docx'},
            "Django views": {
                "query": f"Write Django web CURD (views.py) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{model_name}.docx',
                "outpath":  f'{base_dir}/{model_name}views.docx'},
            "Django api views": {
                "query": f"Write Django api CURD (views.py) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{model_name}.docx',
                "outpath":  f'{base_dir}/{model_name}apiviews.docx'},
            "Django URLs": {
                "query": f"Write Django (urls.py) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{model_name}views.docx',
                "outpath":  f'{base_dir}/{model_name}urls.docx'},
            "Django apis": {
                "query": f"Write Django (urls.py) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{model_name}apiviews.docx',
                "outpath":  f'{base_dir}/{model_name}urls.docx'},
            "Django template": {
                "query":f"Write Django template fragments based code templates (HTML) code for the model '{model_name}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.Include appropriate HTML structure, template tags, and any necessary style.css and script.js.",
                "inputpath": f'{base_dir}/{model_name}views.docx',
                "outpath":  f'{base_dir}/{model_name}html.docx'},
        }
        project_info = {}
        project_details = ""
        os.makedirs(base_dir, exist_ok=True)
        for section, details in sections_queries.items():
            print("query", details.get('query'))
            #query = f"For Project {projectname}, {details.get('query')}"
            response = self.generate_response_from_document(details.get("inputpath"), details.get('query'), project_info)
            #if section == "Django Directory":
            #    self.create_directory_structure_from_text(response, base_dir)
            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key] = response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)

            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path = details.get("outpath")
            # Save the document
            doc.save(path)
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project Business analysis Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"SRS document saved to {output_path}"
        return "Project Code has been completed."
    def extract_model_names(self,text):
        pattern = r'class (\w+)\(models.Model\):'
        matches = re.findall(pattern, text)
        unique_model_names = list(set(matches))
        if not unique_model_names:
            pattern = r'class (\w+)\(db.Model\):'
            matches = re.findall(pattern, text)
            unique_model_names = list(set(matches))
        print("Extracted Model Names:")
        print(unique_model_names)
        return unique_model_names
        
    def extract_table_names(self,docx_path):
        doc = Document(docx_path)
        sql_text = '\n'.join([para.text for para in doc.paragraphs])
        # Regex pattern to find CREATE TABLE statements
        pattern = r'CREATE\s+TABLE\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*\('
        
        # Find all matches in the provided SQL text
        matches = re.findall(pattern, sql_text)
        
        # Remove duplicates by converting the list to a set and back to a list
        unique_table_names = list(set(matches))
        
        return unique_table_names
    def software_developer_agent_model(self, file_path, output_path, base_dir,projectid):
        sections_queries = {
        "Django model": {
            "query": f"Generate Django (models.py) code for the database tables based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
            "path": file_path}}
        project_info = {}
        project_details = ""
        for section, details in sections_queries.items():
            print("query", details.get('query'))
            #query = f"For Project {projectname}, {details.get('query')}"
            response = self.generate_response_from_document(details.get("path"), details.get('query'), project_info)
            #if section == "Django Directory":
            #    self.create_directory_structure_from_text(response, base_dir)
            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key] = response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)

            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path = f'Knowldgebase/SD/{section_key}{projectid}.docx'
            # Save the document
            doc.save(path)
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document
        doc = Document()
        doc.add_heading('Django Models Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"SRS document saved to {output_path}"
        return response
        
    def extract_json_from_docx(self,file_path):
        # Load the document
        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"Error opening file: {e}")
            return None
    
        # Extract text from each paragraph in the document
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
    
        # Combine the text into a single string
        combined_text = "\n".join(full_text)
    
        # Regular expression to find JSON data
        json_pattern = r'\[\s*\{[\s\S]*?\}\s*\]'
    
        # Search for the JSON string in the combined text
        matches = re.search(json_pattern, combined_text)
    
        # If a match is found, convert the JSON string to a Python dictionary
        if matches:
            json_data = matches.group(0)
            try:
                data_list = json.loads(json_data)
                # Convert list of dictionaries into a dictionary with 'name' as the key
                data_dict = {item['name']: item for item in data_list}
                if not data_dict:
                    data_dict = {item["name"]: item for item in data_list}
                print(data_dict)
                return data_dict
            except json.JSONDecodeError:
                print("Error decoding JSON")
        else:
            print("No JSON data found")

    def software_developer_agent_perusecase(self, file_path, output_path, base_dir,projectid,usecase,details):
        details=str(details)
        sections_queries = {
            "Django views": {
                "query": f"Write Django web CURD (views.py) code for the usecase: '{usecase}'details:'{details}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": file_path,
                "outpath":  f'{base_dir}/{usecase}views.docx'},
            "Django api views": {
                "query": f"Write Django api CURD (views.py) code for the usecase: '{usecase}'details:'{details}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": file_path,
                "outpath":  f'{base_dir}/{usecase}apiviews.docx'},
            "Django URLs": {
                "query": f"Write Django (urls.py) code for the usecase: '{usecase}'details:'{details}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{usecase}views.docx',
                "outpath":  f'{base_dir}/{usecase}urls.docx'},
            "Django apis": {
                "query": f"Write Django (urls.py) code for the usecase: '{usecase}'details:'{details}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
                "inputpath": f'{base_dir}/{usecase}apiviews.docx',
                "outpath":  f'{base_dir}/{usecase}urls.docx'},
            "Django template": {
                "query":f"Write Django Django template fragments based code templates (HTML) code for the usecase: '{usecase}'details:'{details}' based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.Include appropriate HTML structure, template tags, and any necessary style.css and script.js.",
                "inputpath": f'{base_dir}/{usecase}views.docx',
                "outpath":  f'{base_dir}/{usecase}html.docx'},
        }
        project_info = {}
        project_details = ""
        os.makedirs(base_dir, exist_ok=True)
        for section, details in sections_queries.items():
            print("query", details.get('query'))
            #query = f"For Project {projectname}, {details.get('query')}"
            response = self.generate_response_from_document(details.get("inputpath"), details.get('query'), project_info)
            #if section == "Django Directory":
            #    self.create_directory_structure_from_text(response, base_dir)
            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key] = response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)

            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path = details.get("outpath")
            # Save the document
            doc.save(path)
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project Code ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"SRS document saved to {output_path}"
        return "Project Code has been completed."
    def software_developer_agent_full(self, file_path, output_path, base_dir,projectid):
        os.makedirs(base_dir, exist_ok=True)
        response = self.software_developer_agent_model(file_path, output_path, base_dir,projectid)
        print('response',response)
        model_names = self.extract_model_names(response)
        #model_names = self.extract_table_names(file_path)
        print('modelnames:',model_names)
        usecasepath = f'Knowldgebase/SA/usecase{projectid}.docx'
        usecases = self.extract_json_from_docx(usecasepath)
        print('usecases:',usecases)
        try:
            print('===============Generating CURD===============')
            for model_name in model_names:
                self.software_developer_agent_permodel(file_path, output_path, base_dir,projectid,model_name)
                # Free all unused cached memory
                torch.cuda.empty_cache()
        
                # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
                torch.cuda.reset_max_memory_allocated()
                torch.cuda.reset_max_memory_cached()
                # Free all unused cached memory
                torch.cuda.empty_cache()
        
                # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
                torch.cuda.reset_max_memory_allocated()
                torch.cuda.reset_max_memory_cached()
                # Free all unused cached memory
                torch.cuda.empty_cache()
        
                # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
                torch.cuda.reset_max_memory_allocated()
                torch.cuda.reset_max_memory_cached()
            print('===============Generating Usecase base Views===============')
            if usecases:
                for name, details in usecases.items():
                    print(f"Name: {name}")
                    print(f"Description: {details['description']}")
                    print("-" * 30)  # Just a separator for clarity
                    codepath = f'Knowldgebase/SD/code_{projectid}.docx'
                    self.software_developer_agent_perusecase(codepath, output_path, base_dir,projectid,name,details)
                    # Free all unused cached memory
                    torch.cuda.empty_cache()
            
                    # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
                    torch.cuda.reset_max_memory_allocated()
                    torch.cuda.reset_max_memory_cached()
                    # Free all unused cached memory
                    torch.cuda.empty_cache()
            
                    # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
                    torch.cuda.reset_max_memory_allocated()
                    torch.cuda.reset_max_memory_cached()
                    # Free all unused cached memory
                    torch.cuda.empty_cache()
            
                    # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
                    torch.cuda.reset_max_memory_allocated()
                    torch.cuda.reset_max_memory_cached()
                return 'Code has been generated.'
            else:
                return "No usecase extracted or available."
        except Exception as e:
            print(f"An error occurred: {e}")
            return 'Failed to generate code.'

    def extract_unique_html_filenames(self,html_content):
        """
        Extracts unique .html filenames from the provided HTML content.
    
        Args:
            html_content (str): The HTML content as a string.
    
        Returns:
            list: A list of unique .html filenames.
        """
        # Regular expression to find .html filenames
        html_files = re.findall(r'href="([^"]+\.html)"', html_content)
    
        # Remove duplicates by converting the list to a set and back to a list
        unique_html_files = list(set(html_files))
    
        return unique_html_fil
        
    def software_developer_agent_template(self, file_path, output_path, base_dir,projectid,ui_data):
        
        query= "Generate the detailed  website HTML code and template for each use case provided in provided SRS document, Ensure that the generated code strictly adheres to the information provided without adding any extra content or features not mentioned in the document. Provide detailed and clear answers, focusing solely on the specified requirements. you should mention html file name. you should nav should have link of other linked pages to that html, should add web based images as example."
        path=file_path
#        sections_queries = {
#          "html": { 
#              "query": f"Write static html single page website with (.html files,style.css,script.js) code for the project based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
#            "path": file_path}}
        project_info = {}
#        project_details = ""
#        for section, details in sections_queries.items():
        print("query", query)
        #query = f"For Project {projectname}, {details.get('query')}"
        response = self.generate_response_from_document(path, query, project_info)
        #if section == "Django Directory":
        #    self.create_directory_structure_from_text(response, base_dir)
        # Ensure the section key is formatted correctly
        section_key = "index"
        #project_info[section_key] = response
        # Create a new .docx document
        doc = Document()
        doc.add_heading(section_key, 0)
        doc.add_paragraph(response)
        path = f'Knowldgebase/SD/{section_key}{projectid}.docx'
        # Save the document
        doc.save(path)
        # Append the response to the project_info with a newline
        #if project_details:
        #    project_details += f"\n{section_key}:\n{response}"
        #else:
        #    project_details = f"{section_key}:\n{response}"
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document
        
        #doc = Document()
        #doc.add_heading('Static Website', 0)

        #for section in sections_queries.keys():
        #doc.add_heading(section.capitalize(), level=1)
        #doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"SRS document saved to {output_path}"
        query= "Generate the detailed static website styles.css code for each html provided in given document. this is the ui requirments.Ensure that the generated code strictly adheres to the information provided without adding any extra content or features not mentioned in the document. Provide detailed and clear answers, focusing solely on the specified requirements."
        path=output_path
#        sections_queries = {
#          "html": { 
#              "query": f"Write static html single page website with (.html files,style.css,script.js) code for the project based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.",
#            "path": file_path}}
        project_info = {}
#        project_details = ""
#        for section, details in sections_queries.items():
        print("query", query)
        #query = f"For Project {projectname}, {details.get('query')}"
        response = self.generate_response_from_document(path, query, project_info)
        #if section == "Django Directory":
        #    self.create_directory_structure_from_text(response, base_dir)
        # Ensure the section key is formatted correctly
        section_key = "styles"
        #project_info[section_key] = response
        # Create a new .docx document
        doc = Document()
        doc.add_heading(section_key, 0)
        doc.add_paragraph(response)
        path = f'Knowldgebase/SD/{section_key}_{projectid}.docx'
        # Save the document
        doc.save(path)
        # Append the response to the project_info with a newline
        #if project_details:
        #    project_details += f"\n{section_key}:\n{response}"
        #else:
        #    project_details = f"{section_key}:\n{response}"
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()
        torch.cuda.reset_max_memory_cached()
        # Free all unused cached memory
        torch.cuda.empty_cache()

        # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
        torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document
        
        #doc = Document()
        #doc.add_heading('Static Website', 0)

        #for section in sections_queries.keys():
        #doc.add_heading(section.capitalize(), level=1)
        #doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(path)
        return "Code has been generated."
        
    def software_developer_agent_perhtml(self, file_path, output_path, base_dir,projectid):
        sections_queries = {
                    "query": f"""Write static html single page website with (index.html,style.css,script.js) code for the project based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.""",

             # f"""Write static html single page website with (index.html,style.css,script.js) code for the project based on the provided document with detailed answers. Do not generate any extra information which the user didn't ask for.""",
            "path": file_path} # htmlpage list , font and font color and color pallete
#            "html": {
#                "query": "Generate Django model class code based on the database models provided in the document. Use Python syntax only. Include appropriate field types, relationships (ForeignKey, OneToOneField, ManyToManyField), indexes, constraints, and any necessary methods.",
#                "path": file_path},
        # }
        project_info = {}
        project_details = ""
        for section, details in sections_queries.items():
            print("query", details.get('query'))
            #query = f"For Project {projectname}, {details.get('query')}"
            response = self.generate_response_from_document(details.get("path"), details.get('query'), project_info)
            #if section == "Django Directory":
            #    self.create_directory_structure_from_text(response, base_dir)
            # Ensure the section key is formatted correctly
            section_key = section.replace(", ", "_").replace(" ", "_")
            project_info[section_key] = response
            # Create a new .docx document
            doc = Document()
            doc.add_heading(section_key, 0)

            for section in sections_queries.keys():
                doc.add_paragraph(response)
            path = f'Knowldgebase/SD/{section_key}{projectid}.docx'
            # Save the document
            doc.save(path)
            # Append the response to the project_info with a newline
            if project_details:
                project_details += f"\n{section_key}:\n{response}"
            else:
                project_details = f"{section_key}:\n{response}"
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()
            torch.cuda.reset_max_memory_cached()
            # Free all unused cached memory
            torch.cuda.empty_cache()

            # Optionally, reset the CUDA memory allocator to prevent potential fragmentation issues
            torch.cuda.reset_max_memory_allocated()

        # Create a new .docx document
        doc = Document()
        doc.add_heading('Project Business analysis Document ', 0)

        for section in sections_queries.keys():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(project_info[section.replace(", ", "_").replace(" ", "_")])

        # Save the document
        doc.save(output_path)
        #return f"SRS document saved to {output_path}"
        return "Project Code has been completed."

    def software_developer_agent_fullstatic(self, file_path, output_path, base_dir,projectid):
        os.makedirs(base_dir, exist_ok=True)
        response = self.software_developer_agent_template(file_path, output_path, base_dir,projectid)
        print('response',response)
        model_names = self.extract_model_names(response)
        print('model_names',model_names)
        try:
            for i, model_name in enumerate(model_names):
                self.software_developer_agent_perhtml(file_path, output_path, base_dir,projectid,model_name)
            return 'Code has been generated.'
        except Exception as e:
            print(f"An error occurred: {e}")
            return 'Failed to generate code.'
    def find_json_structure(self, response):
        try:
            # Attempt to find the JSON structure in the response
            json_start = response.index('{')
            json_end = response.rindex('}') + 1
            json_data = response[json_start:json_end]
            return json_data
        except ValueError:
            raise ValueError("No valid JSON structure found in the response")

    def create_directories(self, base_dir, structure):
        for directory in structure['directories']:
            dir_path = os.path.join(base_dir, directory['name'])
            os.makedirs(dir_path, exist_ok=True)

            if 'files' in directory:
                for file_info in directory['files']:
                    if 'contents' in file_info:
                        sub_dir_path = os.path.join(dir_path, file_info['name'])
                        os.makedirs(sub_dir_path, exist_ok=True)
                        for content in file_info['contents']:
                            file_path = os.path.join(sub_dir_path, content)
                            with open(file_path, 'w') as file:
                                pass  # Create an empty file
                    else:
                        file_path = os.path.join(dir_path, file_info['name'])
                        os.makedirs(os.path.dirname(file_path), exist_ok=True)
                        with open(file_path, 'w') as file:
                            pass  # Create an empty file

    def generate_project_structure_from_response(self, response, base_dir):
        json_data = self.find_json_structure(response)
        structure = json.loads(json_data)
        self.create_directories(base_dir, structure)

    def create_directory_structure_from_text(self, response, base_dir):
        lines = response.strip().split('\n')
        current_path = base_dir
        path_stack = []

        for line in lines:
            # Ignore empty lines
            if not line.strip():
                continue

            # Calculate the level of indentation
            indent_level = len(re.match(r"^\s*\*\s*", line).group()) - 2

            # Extract the name and type of the file/directory
            name_type_match = re.match(r"^\s*-\s*(.*) \((.*)\)", line.strip())
            if name_type_match:
                name, file_type = name_type_match.groups()
            else:
                name = line.strip().split(' ')[-1]
                file_type = "directory"

            # Calculate the current depth
            while path_stack and indent_level <= path_stack[-1][1]:
                path_stack.pop()

            # Get the directory/file name
            current_path = os.path.join(path_stack[-1][0], name) if path_stack else os.path.join(base_dir, name)

            # Create directory or file
            if file_type == "directory":
                os.makedirs(current_path, exist_ok=True)
            else:
                os.makedirs(os.path.dirname(current_path), exist_ok=True)
                with open(current_path, 'w') as file:
                    pass  # Create an empty file

            # Update the stack
            path_stack.append((current_path, indent_level))